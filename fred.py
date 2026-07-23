# -- coding: utf-8 --
import os
import json
import time
import asyncio
import tempfile
import webbrowser
import datetime
import threading
import queue
import math
import urllib.parse
import tkinter as tk

import speech_recognition as sr
import edge_tts
import pygame
import requests
import google.generativeai as genai
from dotenv import load_dotenv
from PIL import ImageGrab, Image, ImageDraw, ImageFilter

try:
    from docx import Document
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False

try:
    import pytesseract
    RUTA_TESSERACT = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.path.exists(RUTA_TESSERACT):
        pytesseract.pytesseract.tesseract_cmd = RUTA_TESSERACT
    OCR_DISPONIBLE = True
except ImportError:
    OCR_DISPONIBLE = False

try:
    import pyperclip
    CLIPBOARD_DISPONIBLE = True
except ImportError:
    CLIPBOARD_DISPONIBLE = False

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
NOMBRE = "Fred"
VARIANTES_NOMBRE = ["fred", "fer", "freddy", "fred,", "fret"]

def detectar_nombre(texto):
    for variante in VARIANTES_NOMBRE:
        if variante in texto:
            return True, texto.replace(variante, "", 1).strip()
    return False, texto.strip()
VOZ = "es-MX-JorgeNeural"
ARCHIVO_MEMORIA = "memoria.json"
ARCHIVO_LECTURAS = "lecturas.json"
MODELO = "gemini-flash-latest"

RUTA_UNITY_HUB = r"C:\Program Files\Unity Hub\Unity Hub.exe"

SEGUNDOS_MODO_CONVERSACION = 45
MINUTOS_CHECKIN = 45

genai.configure(api_key=GEMINI_API_KEY)

atento_hasta = 0.0
ultima_interaccion = time.time()

def cargar_memoria():
    if os.path.exists(ARCHIVO_MEMORIA):
        with open(ARCHIVO_MEMORIA, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"sobre_ti": []}

def guardar_memoria(mem):
    with open(ARCHIVO_MEMORIA, "w", encoding="utf-8") as f:
        json.dump(mem, f, ensure_ascii=False, indent=2)

def cargar_lecturas():
    if os.path.exists(ARCHIVO_LECTURAS):
        with open(ARCHIVO_LECTURAS, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def guardar_lecturas(lecturas):
    with open(ARCHIVO_LECTURAS, "w", encoding="utf-8") as f:
        json.dump(lecturas, f, ensure_ascii=False, indent=2)

memoria = cargar_memoria()

LECTURAS_SESION = cargar_lecturas()

PERSONALIDAD = f"""
Eres {NOMBRE}, el asistente personal por voz del señor.
Hablas SIEMPRE en español, con un tono cálido, cercano y elegante.
Te diriges a él como "señor", con el respeto de un mayordomo,
pero lo quieres como si fuera tu hijo: eres protector, lo animas,
te alegras por él y le das consejos sinceros y honestos.

Eres un compañero de conversación completo: puedes platicar de cualquier
tema (cómo le fue en el día, sus planes, sus dudas), responder preguntas
de cultura general, resolver operaciones matemáticas simples, dar consejos,
opinar, contar datos curiosos, o simplemente charlar sin motivo. NUNCA
digas que no puedes responder algo o que no tienes esa información si es
una pregunta razonable (matemáticas, trivia, opiniones, charla casual);
intenta siempre dar una respuesta útil y natural, como lo haría una
persona culta y cercana.

Respondes CORTO y natural, como en una plática hablada (1 a 3 frases),
porque tus respuestas se van a leer en voz alta.
No uses emojis ni formato, solo texto para hablar.
"""

def contexto_memoria():
    datos = "\n".join(memoria["sobre_ti"]) or "Aún no sé mucho de él."
    return f"Cosas que sé del señor:\n{datos}"

SAFETY_SETTINGS = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_ONLY_HIGH"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_ONLY_HIGH"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_ONLY_HIGH"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_ONLY_HIGH"},
]

modelo = genai.GenerativeModel(MODELO, system_instruction=PERSONALIDAD,
                                safety_settings=SAFETY_SETTINGS)
chat = modelo.start_chat(history=[])

PERSONALIDAD_MAESTRO = f"""
Eres {NOMBRE}, el asistente personal del señor, pero ahora en modo maestro:
te está pidiendo que le expliques o le enseñes algo a fondo. Habla en español,
con el mismo cariño y respeto de mayordomo que lo llama "señor", pero AQUÍ
NO respondas corto: da una explicación completa, clara, bien desarrollada,
con el detalle que el tema merezca (varios párrafos si hace falta). No te
detengas a la mitad ni resumas de más: el señor quiere entender el tema
completo, no solo un resumen de una línea. Organiza la explicación de forma
que se entienda bien al leerse en voz alta (evita usar formato como
asteriscos, títulos o viñetas, ya que esto se convierte a audio).
"""

modelo_maestro = genai.GenerativeModel(MODELO, system_instruction=PERSONALIDAD_MAESTRO,
                                        safety_settings=SAFETY_SETTINGS)

def generar_icono_fred():
    """Genera un ícono tipo orbe brillante para la ventana y la barra de tareas."""
    ruta = os.path.join(tempfile.gettempdir(), "fred_icono.ico")
    tam = 256
    centro = tam // 2
    radio = 78

    img = Image.new("RGBA", (tam, tam), (0, 0, 0, 0))

    # Capa de resplandor (glow) detrás del orbe
    resplandor = Image.new("RGBA", (tam, tam), (0, 0, 0, 0))
    dr = ImageDraw.Draw(resplandor)
    dr.ellipse([centro - radio - 30, centro - radio - 30,
                centro + radio + 30, centro + radio + 30],
               fill=(51, 204, 255, 160))
    resplandor = resplandor.filter(ImageFilter.GaussianBlur(22))
    img = Image.alpha_composite(img, resplandor)

    # Orbe principal con degradado simple
    orbe = Image.new("RGBA", (tam, tam), (0, 0, 0, 0))
    do = ImageDraw.Draw(orbe)
    for i in range(radio, 0, -1):
        t = i / radio
        color = (
            int(20 + (80 - 20) * (1 - t)),
            int(60 + (190 - 60) * (1 - t)),
            int(110 + (255 - 110) * (1 - t)),
            255,
        )
        do.ellipse([centro - i, centro - i, centro + i, centro + i], fill=color)
    do.ellipse([centro - radio, centro - radio, centro + radio, centro + radio],
                outline=(140, 230, 255, 255), width=6)
    img = Image.alpha_composite(img, orbe)

    img.save(ruta, sizes=[(16, 16), (32, 32), (48, 48), (128, 128), (256, 256)])
    return ruta

class VentanaFred:
    def __init__(self):
        self.ventana = tk.Tk()
        self.ventana.title("Fred")
        try:
            ruta_icono = generar_icono_fred()
            self.ventana.iconbitmap(ruta_icono)
        except Exception as e:
            print("No se pudo poner el ícono personalizado:", e)
        self.ventana.geometry("340x540")
        self.ventana.minsize(300, 460)
        self.ventana.configure(bg="#0a0a12")

        # --- Zona del orbe (compacta) ---
        self.canvas = tk.Canvas(self.ventana, width=340, height=170,
                                 bg="#0a0a12", highlightthickness=0)
        self.canvas.pack(fill="x")

        self.cx, self.cy = 170, 78
        self.radio_base = 42

        # Anillos de resplandor detrás del orbe
        self.glow2 = self.canvas.create_oval(0, 0, 0, 0, fill="#0e2438", outline="")
        self.glow1 = self.canvas.create_oval(0, 0, 0, 0, fill="#123049", outline="")
        self.orbe = self.canvas.create_oval(0, 0, 0, 0, fill="#112233",
                                             outline="#33aaff", width=2)
        self.nucleo = self.canvas.create_oval(0, 0, 0, 0, fill="#7fd8ff", outline="")

        self.texto_estado = self.canvas.create_text(
            170, 152, text="En espera...", fill="#55ccff",
            font=("Segoe UI", 10)
        )

        self.hablando = False
        self.microfono_activo = True
        self.atento = False
        self.fase = 0

        self.canvas.bind("<Button-1>", self._alternar_microfono)

        # --- Zona de chat ---
        marco_chat = tk.Frame(self.ventana, bg="#0a0a12")
        marco_chat.pack(fill="both", expand=True, padx=10, pady=(0, 6))

        self.chat = tk.Text(marco_chat, bg="#11121c", fg="#d8e6f0",
                             relief="flat", wrap="word", state="disabled",
                             font=("Segoe UI", 10), padx=10, pady=8,
                             insertbackground="#d8e6f0")
        scroll = tk.Scrollbar(marco_chat, command=self.chat.yview,
                               bg="#11121c", troughcolor="#0a0a12", width=8)
        self.chat.configure(yscrollcommand=scroll.set)
        scroll.pack(side="right", fill="y")
        self.chat.pack(side="left", fill="both", expand=True)

        self.chat.tag_configure("tu", foreground="#8fd48f",
                                 font=("Segoe UI", 10, "bold"))
        self.chat.tag_configure("fred", foreground="#55ccff",
                                 font=("Segoe UI", 10, "bold"))
        self.chat.tag_configure("aviso", foreground="#888899",
                                 font=("Segoe UI", 9, "italic"))

        # --- Cajita para escribir ---
        marco_entrada = tk.Frame(self.ventana, bg="#0a0a12")
        marco_entrada.pack(fill="x", padx=10, pady=(0, 10))

        self.PLACEHOLDER = "¿Alguna duda?"
        self.entrada = tk.Entry(marco_entrada, bg="#1a1c2a", fg="#666677",
                                 insertbackground="#d8e6f0", relief="flat",
                                 font=("Segoe UI", 10))
        self.entrada.pack(fill="x", ipady=7)
        self.entrada.insert(0, self.PLACEHOLDER)
        self.entrada.bind("<FocusIn>", self._quitar_placeholder)
        self.entrada.bind("<FocusOut>", self._poner_placeholder)
        self.entrada.bind("<Return>", self._enviar_texto)

        self._animar()

    def _quitar_placeholder(self, event=None):
        if self.entrada.get() == self.PLACEHOLDER:
            self.entrada.delete(0, tk.END)
            self.entrada.configure(fg="#d8e6f0")

    def _poner_placeholder(self, event=None):
        if not self.entrada.get().strip():
            self.entrada.delete(0, tk.END)
            self.entrada.insert(0, self.PLACEHOLDER)
            self.entrada.configure(fg="#666677")

    def agregar_chat(self, quien, texto):
        def _hacer():
            self.chat.configure(state="normal")
            if quien == "aviso":
                self.chat.insert(tk.END, f"{texto}\n", "aviso")
            else:
                etiqueta = "Tú" if quien == "tu" else NOMBRE
                self.chat.insert(tk.END, f"{etiqueta}: ", quien)
                self.chat.insert(tk.END, f"{texto}\n\n")
            self.chat.configure(state="disabled")
            self.chat.see(tk.END)
        try:
            self.ventana.after(0, _hacer)
        except Exception:
            pass

    def _enviar_texto(self, event=None):
        texto = self.entrada.get().strip()
        if not texto or texto == self.PLACEHOLDER:
            return
        self.entrada.delete(0, tk.END)
        self.agregar_chat("tu", texto)
        cmd = texto.replace(NOMBRE.lower(), "").strip() if NOMBRE.lower() in texto.lower() else texto
        if self.hablando:
            self.agregar_chat("aviso", "Anotado. Fred le responde en cuanto termine de hablar...")
            cola_preguntas_texto.put(cmd)
        else:
            threading.Thread(target=procesar_comando, args=(cmd,), daemon=True).start()

    def _alternar_microfono(self, event=None):
        self.microfono_activo = not self.microfono_activo
        if self.microfono_activo:
            self.canvas.itemconfig(self.texto_estado, text="Escuchando...")
        else:
            self.canvas.itemconfig(self.texto_estado, text="Micrófono en pausa (clic para reanudar)")

    def _coords_circulo(self, item, r):
        self.canvas.coords(item, self.cx - r, self.cy - r, self.cx + r, self.cy + r)

    def _animar(self):
        if self.hablando:
            self.fase += 0.25
            pulso = int(12 * (1 + math.sin(self.fase)))
            r = self.radio_base + pulso
            color, borde = "#1d5f8a", "#55ddff"
        elif not self.microfono_activo:
            r = self.radio_base
            color, borde = "#4a1a1a", "#ff5555"
        elif self.atento:
            self.fase += 0.08
            pulso = int(4 * (1 + math.sin(self.fase)))
            r = self.radio_base + pulso
            color, borde = "#173a25", "#4fd47f"
        else:
            self.fase += 0.04
            pulso = int(2 * (1 + math.sin(self.fase)))
            r = self.radio_base + pulso
            color, borde = "#112233", "#33aaff"

        self._coords_circulo(self.glow2, r + 26)
        self._coords_circulo(self.glow1, r + 13)
        self._coords_circulo(self.orbe, r)
        nucleo_r = max(6, int(r * 0.22 + (3 * abs(math.sin(self.fase * 1.5)))))
        self._coords_circulo(self.nucleo, nucleo_r)
        self.canvas.itemconfig(self.orbe, fill=color, outline=borde)
        self.ventana.after(40, self._animar)

    def set_hablando(self, activo, texto=None):
        self.hablando = activo
        estado = "Hablando..." if activo else "En espera..."
        if texto:
            estado = texto
        self.canvas.itemconfig(self.texto_estado, text=estado)

    def iniciar(self):
        self.ventana.mainloop()

class VentanaEnsayo:
    """Ventana aparte donde Fred escribe el ensayo y recibe sugerencias de mejora."""
    def __init__(self, padre):
        self.win = tk.Toplevel(padre)
        self.win.title("Lienzo — Fred")
        self.win.geometry("560x640")
        self.win.configure(bg="#0a0a12")

        titulo = tk.Label(self.win, text="Lienzo de Fred", bg="#0a0a12",
                           fg="#55ccff", font=("Segoe UI", 13, "bold"))
        titulo.pack(pady=(10, 4))

        marco = tk.Frame(self.win, bg="#0a0a12")
        marco.pack(fill="both", expand=True, padx=12, pady=(0, 6))

        self.texto = tk.Text(marco, bg="#11121c", fg="#e8eef4", relief="flat",
                              wrap="word", font=("Georgia", 11), padx=14, pady=12,
                              insertbackground="#e8eef4")
        scroll = tk.Scrollbar(marco, command=self.texto.yview,
                               bg="#11121c", troughcolor="#0a0a12", width=8)
        self.texto.configure(yscrollcommand=scroll.set)
        scroll.pack(side="right", fill="y")
        self.texto.pack(side="left", fill="both", expand=True)

        self.estado = tk.Label(self.win, text="", bg="#0a0a12", fg="#888899",
                                font=("Segoe UI", 9, "italic"))
        self.estado.pack()

        marco_sug = tk.Frame(self.win, bg="#0a0a12")
        marco_sug.pack(fill="x", padx=12, pady=(2, 12))

        self.PLACEHOLDER = "Escriba una mejora o sugerencia y presione Enter..."
        self.sugerencia = tk.Entry(marco_sug, bg="#1a1c2a", fg="#666677",
                                    insertbackground="#e8eef4", relief="flat",
                                    font=("Segoe UI", 10))
        self.sugerencia.pack(fill="x", ipady=7)
        self.sugerencia.insert(0, self.PLACEHOLDER)
        self.sugerencia.bind("<FocusIn>", self._quitar_ph)
        self.sugerencia.bind("<FocusOut>", self._poner_ph)
        self.sugerencia.bind("<Return>", self._enviar_sugerencia)

        self.ensayo_actual = ""
        self.ocupado = False

    def _quitar_ph(self, event=None):
        if self.sugerencia.get() == self.PLACEHOLDER:
            self.sugerencia.delete(0, tk.END)
            self.sugerencia.configure(fg="#e8eef4")

    def _poner_ph(self, event=None):
        if not self.sugerencia.get().strip():
            self.sugerencia.delete(0, tk.END)
            self.sugerencia.insert(0, self.PLACEHOLDER)
            self.sugerencia.configure(fg="#666677")

    def poner_estado(self, msg):
        try:
            self.win.after(0, lambda: self.estado.configure(text=msg))
        except Exception:
            pass

    def escribir_ensayo(self, texto_completo):
        """Escribe el ensayo con efecto de tecleo progresivo."""
        self.ensayo_actual = texto_completo
        def _escribir():
            self.texto.delete("1.0", tk.END)
            self._tecleo(texto_completo, 0)
        try:
            self.win.after(0, _escribir)
        except Exception:
            pass

    def _tecleo(self, texto_completo, i):
        pedazo = 8
        if i < len(texto_completo):
            self.texto.insert(tk.END, texto_completo[i:i + pedazo])
            self.texto.see(tk.END)
            self.win.after(12, lambda: self._tecleo(texto_completo, i + pedazo))
        else:
            self.poner_estado("Listo. Puede pedir mejoras abajo.")
            self.ocupado = False

    def _enviar_sugerencia(self, event=None):
        pedido = self.sugerencia.get().strip()
        if not pedido or pedido == self.PLACEHOLDER or self.ocupado:
            return
        self.sugerencia.delete(0, tk.END)
        self.ocupado = True
        self.poner_estado("Fred está aplicando su sugerencia...")
        threading.Thread(target=mejorar_ensayo, args=(pedido,), daemon=True).start()

ventana_ensayo = None

ui = VentanaFred()

cola_preguntas_texto = queue.Queue()

pygame.mixer.init()

async def _generar(texto, ruta):
    await edge_tts.Communicate(texto, VOZ).save(ruta)

def hablar(texto):
    print(f"{NOMBRE}: {texto}")
    ui.agregar_chat("fred", texto)
    ui.set_hablando(True, "Hablando...")
    ruta = os.path.join(tempfile.gettempdir(), "fred.mp3")
    asyncio.run(_generar(texto, ruta))
    pygame.mixer.music.load(ruta)
    pygame.mixer.music.play()
    while pygame.mixer.music.get_busy():
        pygame.time.Clock().tick(10)
    pygame.mixer.music.unload()
    ui.set_hablando(False, "En espera...")

reconocedor = sr.Recognizer()
mic = sr.Microphone()

_turnos_desde_calibracion = [999]  # fuerza calibración la primera vez

def escuchar():
    if not ui.microfono_activo:
        pygame.time.wait(300)
        return ""

    if ui.atento:
        ui.set_hablando(False, "Atento, señor...")
    else:
        ui.set_hablando(False, "Escuchando...")

    with mic as fuente:
        if _turnos_desde_calibracion[0] >= 20:
            reconocedor.adjust_for_ambient_noise(fuente, duration=0.3)
            reconocedor.energy_threshold = max(reconocedor.energy_threshold, 300)
            _turnos_desde_calibracion[0] = 0
        else:
            _turnos_desde_calibracion[0] += 1
        reconocedor.pause_threshold = 1.0
        try:
            audio = reconocedor.listen(fuente, timeout=5, phrase_time_limit=12)
        except sr.WaitTimeoutError:
            return ""
        except Exception:
            return ""

    if not ui.microfono_activo:
        # Se pausó justo mientras capturaba audio: descartar por completo.
        return ""

    ui.set_hablando(False, "Pensando...")
    try:
        return reconocedor.recognize_google(audio, language="es-MX").lower()
    except sr.UnknownValueError:
        try:
            return reconocedor.recognize_google(audio, language="en-US").lower()
        except Exception:
            return ""
    except Exception:
        return ""

def tomar_pantalla():
    captura = ImageGrab.grab()
    return captura

def extraer_texto_pantalla(imagen):
    if not OCR_DISPONIBLE:
        return None
    try:
        texto = pytesseract.image_to_string(imagen, lang="spa+eng")
        texto = texto.strip()
        return texto if len(texto) > 20 else None
    except Exception as e:
        print("Error OCR:", e)
        return None

def analizar_pantalla(pregunta, modo="breve"):
    """modo: 'breve' (solo imagen, rápido), 'resumen' (OCR + corto pero completo),
    'detallado' (OCR + explicación larga)."""
    try:
        imagen = tomar_pantalla()
        texto_extraido = extraer_texto_pantalla(imagen) if modo in ("resumen", "detallado") else None

        if modo == "detallado":
            if texto_extraido:
                prompt = (f"{contexto_memoria()}\n\nEste es el texto exacto que se extrajo "
                          f"de la pantalla del señor mediante OCR:\n\n---\n{texto_extraido}\n---\n\n"
                          f"El señor pide: {pregunta}\n\nUsa ESTE texto extraído como la fuente "
                          f"principal y más confiable (es más preciso que lo que se ve en la "
                          f"imagen). Ignora fragmentos de OCR que parezcan basura o menús de "
                          f"navegador sin relación con el contenido principal.")
            else:
                prompt = f"{contexto_memoria()}\n\nEl señor te muestra su pantalla y pide: {pregunta}"
            respuesta = modelo_maestro.generate_content([prompt, imagen])

        elif modo == "resumen":
            fuente_texto = (f"Este es el texto exacto extraído por OCR de la pantalla:\n\n"
                             f"---\n{texto_extraido}\n---\n\n" if texto_extraido else "")
            prompt = (f"{contexto_memoria()}\n\n{fuente_texto}"
                      f"El señor pide: {pregunta}\n\nLee y considera TODO el texto/contenido "
                      f"completo (no te bases solo en el título o lo primero que veas), pero "
                      f"da tu respuesta CORTA: un resumen de 2 a 4 frases con solo los puntos "
                      f"más importantes, hablado y natural, como su mayordomo Fred.")
            respuesta = modelo.generate_content([prompt, imagen])

        else:
            prompt = (f"{contexto_memoria()}\n\nEl señor te muestra su pantalla y "
                      f"pregunta: {pregunta}. Describe o ayuda de forma breve y "
                      f"hablada, como su mayordomo Fred, en 1 a 3 frases.")
            respuesta = modelo.generate_content([prompt, imagen])

        hablar(respuesta.text.strip())

        if modo in ("resumen", "detallado"):
            LECTURAS_SESION.append({
                "hora": datetime.datetime.now().strftime("%d/%m %H:%M"),
                "pregunta": pregunta,
                "respuesta": respuesta.text.strip(),
            })
            guardar_lecturas(LECTURAS_SESION)
            chat.history.append({"role": "user",
                                  "parts": [f"(Le mostré mi pantalla y pedí: {pregunta})"]})
            chat.history.append({"role": "model", "parts": [respuesta.text.strip()]})
    except Exception as e:
        error_str = str(e)
        print("Error vision:", error_str)
        if "429" in error_str:
            hablar("Señor, alcancé mi límite de consultas con la inteligencia artificial por hoy.")
        else:
            hablar("Disculpe, señor, no pude analizar la pantalla en este momento.")

def crear_ensayo(cmd):
    global ventana_ensayo
    if not LECTURAS_SESION:
        hablar("Señor, todavía no he leído nada con qué construir el texto.")
        return

    hablar("Con mucho gusto, señor. Abriendo el lienzo.")
    try:
        def _abrir():
            global ventana_ensayo
            ventana_ensayo = VentanaEnsayo(ui.ventana)
            ventana_ensayo.poner_estado("Fred está redactando...")
        ui.ventana.after(0, _abrir)
        time.sleep(0.6)

        bloques = "\n\n".join(
            f"Lectura {i+1} ({l['hora']}): {l['respuesta']}"
            for i, l in enumerate(LECTURAS_SESION)
        )
        prompt = (f"Estas son las lecturas y explicaciones que le diste al señor sobre lo que "
                  f"vio en su pantalla:\n\n{bloques}\n\nEl señor pide: {cmd}\n\n"
                  f"Redacta EXACTAMENTE el tipo de texto que el señor pidió (puede ser un "
                  f"ensayo, un texto informativo, un artículo, una carta, un guion, etc.), "
                  f"EN ESPAÑOL, completo y bien estructurado según lo que corresponda a ese "
                  f"tipo de texto, basándote en el contenido leído. Si las lecturas son "
                  f"consejos de redacción, aplícalos. Escribe SOLO el texto final, sin "
                  f"saludos ni comentarios de mayordomo, sin markdown, en párrafos normales.")
        respuesta = modelo_maestro.generate_content(prompt)
        texto_ensayo = respuesta.text.strip()

        if ventana_ensayo:
            ventana_ensayo.escribir_ensayo(texto_ensayo)
        hablar("Texto listo, señor. Puede escribir mejoras en la parte de abajo de la ventana.")
    except Exception as e:
        error_str = str(e)
        print("Error creando ensayo:", error_str)
        if "429" in error_str:
            hablar("Señor, alcancé mi límite de consultas con la inteligencia artificial por hoy.")
        else:
            hablar("Disculpe, señor, tuve un problema al redactar el ensayo.")

def mejorar_ensayo(pedido):
    global ventana_ensayo
    if not ventana_ensayo or not ventana_ensayo.ensayo_actual:
        return
    try:
        prompt = (f"Este es el texto actual:\n\n---\n{ventana_ensayo.ensayo_actual}\n---\n\n"
                  f"El señor pide esta mejora: {pedido}\n\n"
                  f"Reescribe el texto COMPLETO aplicando esa mejora, manteniendo lo que ya "
                  f"está bien y respetando el tipo de texto que es. Devuelve solo el texto "
                  f"final, sin comentarios ni markdown.")
        respuesta = modelo_maestro.generate_content(prompt)
        ventana_ensayo.escribir_ensayo(respuesta.text.strip())
    except Exception as e:
        print("Error mejorando ensayo:", e)
        ventana_ensayo.poner_estado("Hubo un problema al aplicar la mejora. Intente de nuevo.")
        ventana_ensayo.ocupado = False

def generar_resumen_documento(cmd_original=""):
    if not LECTURAS_SESION:
        hablar("Señor, todavía no he leído nada que pueda resumir.")
        return

    quiere_word = "word" in cmd_original or "documento de word" in cmd_original

    hablar("Enseguida, señor. Preparando el resumen.")
    try:
        bloques = "\n\n".join(
            f"Lectura {i+1} ({l['hora']}) — pregunta: {l['pregunta']}\nExplicación dada: {l['respuesta']}"
            for i, l in enumerate(LECTURAS_SESION)
        )
        prompt = (f"A continuación tienes {len(LECTURAS_SESION)} lecturas de pantalla que le "
                  f"hiciste al señor:\n\n{bloques}\n\n"
                  f"Redacta un resumen organizado en español para un documento: "
                  f"identifica los puntos más importantes de TODAS las lecturas en conjunto, "
                  f"agrupa ideas relacionadas si las hay, y sé claro y completo. Escribe solo "
                  f"el contenido del resumen (sin saludos de mayordomo, esto es para un documento "
                  f"formal, no para hablarlo en voz alta). Usa párrafos normales, sin markdown.")
        respuesta = modelo_maestro.generate_content(prompt)
        texto_resumen = respuesta.text.strip()

        if quiere_word and DOCX_DISPONIBLE:
            carpeta = os.path.join(os.path.expanduser("~"), "Documents", "Resúmenes de Fred")
            os.makedirs(carpeta, exist_ok=True)
            nombre_archivo = f"Resumen {datetime.datetime.now().strftime('%Y-%m-%d %H-%M')}.docx"
            ruta_completa = os.path.join(carpeta, nombre_archivo)

            doc = Document()
            doc.add_heading("Resumen de lecturas — Fred", level=1)
            doc.add_paragraph(f"Generado el {datetime.datetime.now().strftime('%d/%m/%Y a las %H:%M')}")
            for parrafo in texto_resumen.split("\n"):
                if parrafo.strip():
                    doc.add_paragraph(parrafo.strip())
            doc.save(ruta_completa)
            os.startfile(ruta_completa)
            hablar("Listo, señor. Ya le abrí el documento de Word con el resumen.")
            return

        if not CLIPBOARD_DISPONIBLE:
            hablar("Señor, me falta instalar una herramienta para copiar el texto. Avíseme para configurarla.")
            return

        pyperclip.copy(texto_resumen)
        webbrowser.open("https://docs.new")
        hablar("Listo, señor. Le copié el resumen y le abrí un Google Doc en blanco. "
               "Solo pegue con Control V cuando cargue la página.")
    except Exception as e:
        print("Error generando documento:", e)
        hablar("Disculpe, señor, tuve un problema al preparar el resumen.")

YOUTUBE_API_KEY = os.getenv("YOUTUBE_API_KEY")

def buscar_video_youtube(consulta):
    """Busca en YouTube y devuelve el ID del primer video, o None si falla/no hay clave."""
    if not YOUTUBE_API_KEY:
        return None
    try:
        resp = requests.get(
            "https://www.googleapis.com/youtube/v3/search",
            params={
                "part": "id",
                "q": consulta,
                "type": "video",
                "maxResults": 1,
                "key": YOUTUBE_API_KEY,
            },
            timeout=5,
        )
        items = resp.json().get("items", [])
        if items:
            return items[0]["id"]["videoId"]
    except Exception as e:
        print("Error buscando en YouTube:", e)
    return None

def ejecutar_accion(cmd):
    if cmd.startswith("pon ") or cmd.startswith("reproduce ") or cmd.startswith("play "):
        cancion = cmd
        for palabra in ["pon ", "reproduce ", "play "]:
            cancion = cancion.replace(palabra, "", 1) if cancion.startswith(palabra) else cancion
        cancion = cancion.strip()
        if cancion:
            video_id = buscar_video_youtube(cancion)
            if video_id:
                webbrowser.open(f"https://www.youtube.com/watch?v={video_id}")
                hablar(f"Reproduciendo {cancion}, señor.")
            else:
                consulta = urllib.parse.quote(cancion)
                webbrowser.open(f"https://www.youtube.com/results?search_query={consulta}")
                hablar(f"Buscando {cancion} en YouTube, señor.")
            return True

    if "modo juego" in cmd:
        try:
            os.startfile("ms-xbox:")
        except Exception:
            pass
        hablar("Modo juego activado, señor. Que tenga una buena partida.")
        return True

    if any(frase in cmd for frase in [
        "lee esto", "léelo", "leelo", "explícame esto", "explicame esto",
        "qué dice aquí", "que dice aqui", "explica esto", "explícame todo",
        "explicame todo", "detalladamente", "a detalle", "a fondo",
    ]):
        hablar("Permítame leer, señor.")
        prompt_lectura = ("Lee con atención todo el texto visible en la pantalla (puede ser "
                          "un artículo, una guía, código, o cualquier documento), y explícaselo "
                          "al señor de forma completa y bien desarrollada: qué dice, qué significa, "
                          "y cualquier punto importante que contenga. No lo resumas en una sola "
                          "línea — desarrolla la explicación con el detalle que el contenido merezca.")
        analizar_pantalla(prompt_lectura, modo="detallado")
        return True

    if any(frase in cmd for frase in [
        "resume qué veo", "resume que veo", "resume lo que veo", "resúmeme esto",
        "resumeme esto", "dame un resumen de esto", "resumen de esto",
        "resúmelo", "resumelo",
    ]):
        analizar_pantalla("Resume lo que ves en la pantalla.", modo="resumen")
        return True

    if any(frase in cmd for frase in [
        "haz un ensayo", "hazme un ensayo", "escribe un ensayo", "crea un ensayo",
        "haz un texto", "hazme un texto", "escribe un texto", "crea un texto",
        "haz un artículo", "haz un articulo", "escribe un artículo", "escribe un articulo",
        "haz una carta", "escribe una carta", "redacta",
    ]):
        threading.Thread(target=crear_ensayo, args=(cmd,), daemon=True).start()
        return True

    if any(frase in cmd for frase in [
        "haz un documento", "documento con el resumen", "resumen de la sesión",
        "resumen de la sesion", "resumen de todo lo que", "haz un resumen de todo",
        "resumen de todo lo que hemos visto", "resumen de todo lo que vimos",
    ]):
        generar_resumen_documento(cmd)
        return True

    if "ayúdame con esto" in cmd or "ayudame con esto" in cmd or "ayúdame con este error" in cmd or "ayudame con este error" in cmd:
        hablar("Déjeme revisar, señor.")
        analizar_pantalla("Ayúdame con lo que ves en mi pantalla, especialmente si hay algún error o problema visible.", modo="detallado")
        return True

    if "qué estoy viendo" in cmd or "que estoy viendo" in cmd or "qué ves" in cmd or "que ves" in cmd:
        hablar("Permítame ver su pantalla, señor.")
        analizar_pantalla("¿Qué estoy viendo en mi pantalla?", modo="breve")
        return True



    if "youtube" in cmd:
        webbrowser.open("https://youtube.com")
        hablar("Enseguida, señor. Abriendo YouTube.")
        return True

    if "spotify" in cmd:
        webbrowser.open("https://open.spotify.com")
        hablar("Abriendo Spotify, señor.")
        return True

    if "busca en google" in cmd or "buscar en google" in cmd:
        consulta = cmd.replace("busca en google", "").replace("buscar en google", "").strip()
        webbrowser.open(f"https://www.google.com/search?q={urllib.parse.quote(consulta)}")
        hablar(f"Buscando {consulta}, señor.")
        return True

    if "google" in cmd:
        webbrowser.open("https://google.com")
        hablar("Listo, señor.")
        return True

    if "qué hora" in cmd or "que hora" in cmd or "what time" in cmd:
        hora = datetime.datetime.now().strftime("%I:%M")
        hablar(f"Son las {hora}, señor.")
        return True

    if "qué día" in cmd or "que dia" in cmd or "fecha" in cmd:
        fecha = datetime.datetime.now().strftime("%d de %B")
        hablar(f"Hoy es {fecha}, señor.")
        return True

    if "abre mis documentos" in cmd or "abre documentos" in cmd:
        os.startfile(os.path.expanduser("~\\Documents"))
        hablar("Abriendo sus documentos, señor.")
        return True

    if "abre el escritorio" in cmd or "abre mi escritorio" in cmd:
        os.startfile(os.path.expanduser("~\\Desktop"))
        hablar("Abriendo su escritorio, señor.")
        return True

    if "abre xbox" in cmd or "abre el xbox" in cmd:
        os.startfile("ms-xbox:")
        hablar("Abriendo la app de Xbox, señor.")
        return True

    if "abre steam" in cmd:
        os.startfile("steam:")
        hablar("Abriendo Steam, señor.")
        return True

    if "abre unity" in cmd:
        try:
            os.startfile(RUTA_UNITY_HUB)
            hablar("Abriendo Unity Hub, señor.")
        except Exception:
            hablar("No encontré Unity Hub en esa ruta, señor. Habría que revisar la ubicación en el código.")
        return True

    if "clima" in cmd or "temperatura" in cmd:
        try:
            resp = requests.get("https://wttr.in/?format=3&lang=es", timeout=6)
            hablar(f"Señor, {resp.text.strip()}")
        except Exception:
            hablar("No pude consultar el clima en este momento, señor.")
        return True

    return False

def responder(cmd):
    prompt = f"{contexto_memoria()}\n\nEl señor dijo: {cmd}"
    texto = None
    intentos = 3
    for intento in range(intentos):
        try:
            texto = chat.send_message(prompt).text.strip()
            break
        except Exception as e:
            error_str = str(e)
            print("Error IA:", error_str)
            if "429" in error_str:
                if "PerDay" in error_str or "per_day" in error_str.lower():
                    texto = "Señor, alcancé mi límite diario de conversaciones con la inteligencia artificial. Las acciones directas siguen a su servicio."
                    break
                if intento < intentos - 1:
                    pygame.time.wait(4000)
                    continue
                texto = "Señor, los servidores están saturados en este momento. Intentemos en un minuto."
                break
            texto = "Disculpe, señor, tuve un problema para pensar eso."
            break
    hablar(texto)

def checkin_periodico():
    global ultima_interaccion
    while True:
        time.sleep(60)
        minutos_sin_hablar = (time.time() - ultima_interaccion) / 60
        if minutos_sin_hablar >= MINUTOS_CHECKIN and ui.microfono_activo and not ui.hablando:
            hablar("¿Necesita algo, señor?")
            ultima_interaccion = time.time()

def procesar_comando(cmd):
    """Procesa un comando ya limpio (sin 'Fred'), venga de voz o de texto escrito."""
    global ultima_interaccion
    ultima_interaccion = time.time()

    if not cmd:
        hablar("Dígame, señor.")
        return True  # seguir activo

    if "adiós" in cmd or "hasta luego" in cmd or "duérmete" in cmd:
        hablar("Que descanse, señor. Aquí estaré.")
        return False  # señal de apagar

    if cmd.startswith("recuerda que"):
        dato = cmd.replace("recuerda que", "").strip()
        memoria["sobre_ti"].append(dato)
        guardar_memoria(memoria)
        hablar("Lo tendré presente, señor.")
        return True

    if not ejecutar_accion(cmd):
        responder(cmd)
    return True

def vigilar_cola_preguntas():
    """Espera a que Fred deje de hablar y procesa preguntas escritas que quedaron en cola."""
    while True:
        cmd = cola_preguntas_texto.get()  # se detiene aquí hasta que llegue algo
        while ui.hablando:
            time.sleep(0.2)
        procesar_comando(cmd)

def bucle_fred():
    global atento_hasta, ultima_interaccion
    hablar(f"A sus órdenes, señor. {NOMBRE} está despierto.")
    while True:
        ui.atento = time.time() < atento_hasta
        texto = escuchar()
        if not texto:
            continue
        print(f"Tú: {texto}")
        ui.agregar_chat("tu", texto)

        en_modo_conversacion = time.time() < atento_hasta
        tiene_nombre, cmd_sin_nombre = detectar_nombre(texto)

        if not tiene_nombre and not en_modo_conversacion:
            print(f'(No detecté "Fred" en: "{texto}" — ignorando)')
            ui.set_hablando(False, 'No escuché "Fred", intente de nuevo')
            time.sleep(1.2)
            continue

        atento_hasta = time.time() + SEGUNDOS_MODO_CONVERSACION

        cmd = cmd_sin_nombre if tiene_nombre else texto.strip()

        if not procesar_comando(cmd):
            break

if __name__ == "__main__":
    hilo = threading.Thread(target=bucle_fred, daemon=True)
    hilo.start()
    hilo_checkin = threading.Thread(target=checkin_periodico, daemon=True)
    hilo_checkin.start()
    hilo_cola = threading.Thread(target=vigilar_cola_preguntas, daemon=True)
    hilo_cola.start()
    ui.iniciar()
